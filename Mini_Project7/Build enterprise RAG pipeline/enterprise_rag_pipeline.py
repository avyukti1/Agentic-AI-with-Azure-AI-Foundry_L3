"""Reusable enterprise RAG pipeline helpers for the notebook demo.

The Streamlit app focuses on chat UI and runtime routing. This module keeps the
notebook workflow importable without executing Streamlit code: load documents,
chunk them, create/update Azure AI Search, embed with Azure OpenAI, retrieve,
answer with grounded citations, and run a lightweight evaluation file.
"""

from __future__ import annotations

import json
import os
import re
from io import BytesIO
from pathlib import Path
from typing import Iterable

import fitz
import openpyxl
import requests
from azure.core.credentials import AzureKeyCredential
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.search.documents.indexes.models import (
    HnswAlgorithmConfiguration,
    HnswParameters,
    SearchableField,
    SearchField,
    SearchFieldDataType,
    SearchIndex,
    SemanticConfiguration,
    SemanticField,
    SemanticPrioritizedFields,
    SemanticSearch,
    SimpleField,
    VectorSearch,
    VectorSearchProfile,
)
from azure.storage.blob import ContainerClient
from docx import Document
from dotenv import load_dotenv
from openai import AzureOpenAI


load_dotenv(dotenv_path=".env", override=True)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".log", ".pdf", ".docx", ".xlsx", ".xlsm"}
DEFAULT_CHUNK_SIZE = 1200
DEFAULT_CHUNK_OVERLAP = 180
SEMANTIC_CONFIG_NAME = "rag-semantic-config"
VECTOR_PROFILE_NAME = "rag-vector-profile"
VECTOR_ALGORITHM_NAME = "rag-hnsw"


def require_env(name: str) -> str:
    """Return a required environment value or raise a readable error."""
    value = os.getenv(name)
    if not value or value.startswith("<"):
        raise RuntimeError(f"Missing required .env value: {name}")
    return value


def get_openai_client() -> AzureOpenAI:
    """Create an Azure OpenAI client using Microsoft Entra authentication."""
    token_provider = get_bearer_token_provider(
        DefaultAzureCredential(),
        "https://cognitiveservices.azure.com/.default",
    )
    return AzureOpenAI(
        azure_endpoint=require_env("AZURE_OPENAI_ENDPOINT"),
        azure_ad_token_provider=token_provider,
        api_version="2024-10-21",
    )


def get_search_client() -> SearchClient:
    """Create an Azure AI Search document client for the configured index."""
    return SearchClient(
        endpoint=require_env("AI_SEARCH_SERVICE_ENDPOINT"),
        index_name=require_env("AI_SEARCH_INDEX_NAME"),
        credential=AzureKeyCredential(require_env("AI_SEARCH_API_KEY")),
    )


def get_index_client() -> SearchIndexClient:
    """Create an Azure AI Search index-management client."""
    return SearchIndexClient(
        endpoint=require_env("AI_SEARCH_SERVICE_ENDPOINT"),
        credential=AzureKeyCredential(require_env("AI_SEARCH_API_KEY")),
    )


def embed_texts(client: AzureOpenAI, texts: list[str]) -> list[list[float]]:
    """Generate embeddings for one or more texts with the configured deployment."""
    if not texts:
        return []
    response = client.embeddings.create(
        model=require_env("EMBEDDING_DEPLOYMENT_NAME"),
        input=texts,
    )
    return [item.embedding for item in response.data]


def parse_document(file_name: str, content: bytes) -> str:
    """Extract text from supported document formats."""
    suffix = Path(file_name).suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".log"}:
        return content.decode("utf-8", errors="ignore")

    if suffix == ".pdf":
        with fitz.open(stream=content, filetype="pdf") as pdf:
            return "\n\n".join(page.get_text("text") for page in pdf)

    if suffix == ".docx":
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)

    if suffix in {".xlsx", ".xlsm"}:
        workbook = openpyxl.load_workbook(BytesIO(content), data_only=True, read_only=True)
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    lines.append(" | ".join(values))
        workbook.close()
        return "\n".join(lines)

    raise ValueError(f"Unsupported document extension: {suffix}")


def split_text(text: str, chunk_size: int = DEFAULT_CHUNK_SIZE, overlap: int = DEFAULT_CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping chunks at paragraph-friendly boundaries."""
    normalized = re.sub(r"\n{3,}", "\n\n", text.strip())
    if not normalized:
        return []

    chunks = []
    start = 0
    while start < len(normalized):
        end = min(start + chunk_size, len(normalized))
        if end < len(normalized):
            boundary = max(normalized.rfind("\n\n", start, end), normalized.rfind(". ", start, end))
            if boundary > start + chunk_size // 2:
                end = boundary + 1
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start = max(end - overlap, 0)
    return chunks


def metadata_from_text(file_name: str, text: str) -> dict[str, str]:
    """Read simple metadata headers and fall back to filename hints."""
    title = Path(file_name).stem.replace("_", " ").replace("-", " ").title()
    department = "General"
    security_level = "internal"

    for line in text.splitlines()[:10]:
        key, _, value = line.partition(":")
        if not value:
            continue
        key = key.strip().lower()
        value = value.strip()
        if key == "title":
            title = value
        elif key == "department":
            department = value
        elif key in {"security", "security_level", "security level"}:
            security_level = value

    lower_name = file_name.lower()
    if department == "General":
        if any(token in lower_name for token in ["hr", "benefit", "employee", "payroll"]):
            department = "HR"
        elif any(token in lower_name for token in ["it", "security", "helpdesk", "vpn"]):
            department = "IT"

    return {"title": title, "department": department, "security_level": security_level}


def make_chunks(file_name: str, content: bytes, source_label: str) -> list[dict]:
    """Parse one file and return Search-ready text chunks without vectors."""
    text = parse_document(file_name, content)
    metadata = metadata_from_text(file_name, text)
    chunks = []
    for index, chunk_text in enumerate(split_text(text), start=1):
        source_file = f"{source_label}: {file_name}" if source_label else file_name
        chunk_key = re.sub(r"[^A-Za-z0-9_-]+", "-", source_file).strip("-").lower()
        chunks.append(
            {
                "id": f"{chunk_key}-{index}",
                "content": chunk_text,
                "title": metadata["title"],
                "source_file": source_file,
                "department": metadata["department"],
                "security_level": metadata["security_level"],
                "chunk_id": str(index),
            }
        )
    return chunks


def load_local_documents() -> list[dict]:
    """Load supported files from RAG_DOCS_DIR."""
    docs_dir = Path(os.getenv("RAG_DOCS_DIR", "./data/documents"))
    if not docs_dir.exists():
        raise RuntimeError(f"Local document folder was not found: {docs_dir}")

    documents = []
    for path in sorted(docs_dir.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            documents.extend(make_chunks(path.name, path.read_bytes(), "Local"))
    return documents


def load_blob_documents() -> list[dict]:
    """Load supported blobs from BLOB_CONTAINER_URL."""
    container_url = require_env("BLOB_CONTAINER_URL")
    credential = None if "sig=" in container_url.lower() else DefaultAzureCredential()
    container = ContainerClient.from_container_url(container_url, credential=credential)
    documents = []
    for blob in container.list_blobs():
        if Path(blob.name).suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        content = container.download_blob(blob.name).readall()
        documents.extend(make_chunks(blob.name, content, "Blob"))
    return documents


def graph_get(url: str, token: str) -> dict:
    """Call Microsoft Graph and return JSON."""
    response = requests.get(url, headers={"Authorization": f"Bearer {token}"}, timeout=60)
    response.raise_for_status()
    return response.json()


def get_sharepoint_token() -> str:
    """Acquire a Microsoft Graph token for SharePoint ingestion."""
    response = requests.post(
        f"https://login.microsoftonline.com/{require_env('AZURE_TENANT_ID')}/oauth2/v2.0/token",
        data={
            "client_id": require_env("SHAREPOINT_CLIENT_ID"),
            "client_secret": require_env("SHAREPOINT_CLIENT_SECRET"),
            "scope": "https://graph.microsoft.com/.default",
            "grant_type": "client_credentials",
        },
        timeout=60,
    )
    response.raise_for_status()
    return response.json()["access_token"]


def load_sharepoint_documents() -> list[dict]:
    """Load supported files from the configured SharePoint document library."""
    token = get_sharepoint_token()
    site = graph_get(
        f"https://graph.microsoft.com/v1.0/sites/{require_env('SHAREPOINT_HOSTNAME')}:{require_env('SHAREPOINT_SITE_PATH')}",
        token,
    )
    drives = graph_get(f"https://graph.microsoft.com/v1.0/sites/{site['id']}/drives", token).get("value", [])
    drive_name = os.getenv("SHAREPOINT_DRIVE_NAME", "Documents")
    drive = next((item for item in drives if item.get("name") == drive_name), None)
    if not drive:
        raise RuntimeError(f"SharePoint drive `{drive_name}` was not found.")

    files = []
    stack = [f"https://graph.microsoft.com/v1.0/drives/{drive['id']}/root/children"]
    while stack:
        payload = graph_get(stack.pop(), token)
        for item in payload.get("value", []):
            if "folder" in item:
                stack.append(f"https://graph.microsoft.com/v1.0/drives/{drive['id']}/items/{item['id']}/children")
            elif "file" in item and Path(item.get("name", "")).suffix.lower() in SUPPORTED_EXTENSIONS:
                files.append(item)
        next_link = payload.get("@odata.nextLink")
        if next_link:
            stack.append(next_link)

    documents = []
    for item in files:
        response = requests.get(
            f"https://graph.microsoft.com/v1.0/drives/{drive['id']}/items/{item['id']}/content",
            headers={"Authorization": f"Bearer {token}"},
            timeout=120,
        )
        response.raise_for_status()
        documents.extend(make_chunks(item["name"], response.content, "SharePoint"))
    return documents


def configured_sources() -> list[str]:
    """Return normalized RAG sources from RAG_SOURCE."""
    raw_sources = [item.strip().lower() for item in os.getenv("RAG_SOURCE", "local").split(",") if item.strip()]
    if "all" in raw_sources:
        return ["sharepoint", "blob", "local"]
    return [source for source in raw_sources if source in {"sharepoint", "blob", "local"}]


def load_source_documents() -> list[dict]:
    """Load chunks from each configured source."""
    loaders = {
        "sharepoint": load_sharepoint_documents,
        "blob": load_blob_documents,
        "local": load_local_documents,
    }
    documents = []
    errors = []
    for source in configured_sources():
        try:
            documents.extend(loaders[source]())
        except Exception as exc:
            errors.append(f"{source}: {exc}")

    if not documents and errors:
        raise RuntimeError("No source documents loaded.\n" + "\n".join(errors))
    if errors:
        print("Some sources could not be loaded:")
        for error in errors:
            print(f"- {error}")
    return documents


def create_or_update_index(vector_dimensions: int | None = None) -> SearchIndex:
    """Create or update the configured Azure AI Search vector index."""
    if vector_dimensions is None:
        client = get_openai_client()
        vector_dimensions = len(embed_texts(client, ["dimension probe"])[0])

    fields = [
        SimpleField(name="id", type=SearchFieldDataType.String, key=True, filterable=True),
        SearchableField(name="content", type=SearchFieldDataType.String),
        SearchableField(name="title", type=SearchFieldDataType.String, filterable=True, sortable=True),
        SearchableField(name="source_file", type=SearchFieldDataType.String, filterable=True, sortable=True),
        SimpleField(name="department", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="security_level", type=SearchFieldDataType.String, filterable=True, facetable=True),
        SimpleField(name="chunk_id", type=SearchFieldDataType.String, filterable=True, sortable=True),
        SearchField(
            name="content_vector",
            type=SearchFieldDataType.Collection(SearchFieldDataType.Single),
            searchable=True,
            vector_search_dimensions=vector_dimensions,
            vector_search_profile_name=VECTOR_PROFILE_NAME,
        ),
    ]
    vector_search = VectorSearch(
        algorithms=[
            HnswAlgorithmConfiguration(
                name=VECTOR_ALGORITHM_NAME,
                parameters=HnswParameters(metric="cosine"),
            )
        ],
        profiles=[
            VectorSearchProfile(
                name=VECTOR_PROFILE_NAME,
                algorithm_configuration_name=VECTOR_ALGORITHM_NAME,
            )
        ],
    )
    semantic_search = SemanticSearch(
        configurations=[
            SemanticConfiguration(
                name=SEMANTIC_CONFIG_NAME,
                prioritized_fields=SemanticPrioritizedFields(
                    title_field=SemanticField(field_name="title"),
                    content_fields=[SemanticField(field_name="content")],
                    keywords_fields=[
                        SemanticField(field_name="source_file"),
                        SemanticField(field_name="department"),
                    ],
                ),
            )
        ]
    )
    index = SearchIndex(
        name=require_env("AI_SEARCH_INDEX_NAME"),
        fields=fields,
        vector_search=vector_search,
        semantic_search=semantic_search,
    )
    return get_index_client().create_or_update_index(index)


def batched(items: list[dict], size: int) -> Iterable[list[dict]]:
    """Yield fixed-size batches."""
    for start in range(0, len(items), size):
        yield items[start : start + size]


def ingest_documents() -> dict:
    """Load, embed, and upload configured source documents to Azure AI Search."""
    openai_client = get_openai_client()
    documents = load_source_documents()
    if not documents:
        raise RuntimeError("No documents were found for ingestion.")

    sample_vector = embed_texts(openai_client, [documents[0]["content"]])[0]
    create_or_update_index(vector_dimensions=len(sample_vector))

    documents[0]["content_vector"] = sample_vector
    remaining = documents[1:]
    for batch in batched(remaining, 16):
        vectors = embed_texts(openai_client, [item["content"] for item in batch])
        for item, vector in zip(batch, vectors, strict=True):
            item["content_vector"] = vector

    search_client = get_search_client()
    uploaded = 0
    for batch in batched(documents, 100):
        result = search_client.upload_documents(documents=batch)
        uploaded += sum(1 for item in result if item.succeeded)

    return {"loaded": len(documents), "uploaded": uploaded, "index": require_env("AI_SEARCH_INDEX_NAME")}


def retrieve(question: str, filter_expression: str | None = None, top: int = 3) -> list[dict]:
    """Retrieve relevant chunks with hybrid keyword and vector search."""
    vector = embed_texts(get_openai_client(), [question])[0]
    results = get_search_client().search(
        search_text=question,
        vector_queries=[
            {
                "kind": "vector",
                "vector": vector,
                "fields": "content_vector",
                "k": top,
            }
        ],
        query_type="semantic",
        semantic_configuration_name=SEMANTIC_CONFIG_NAME,
        filter=filter_expression,
        top=top,
        select=["content", "title", "source_file", "department", "security_level", "chunk_id"],
    )
    return [dict(item) for item in results]


def answer_question(question: str, filter_expression: str | None = None, top: int = 3) -> str:
    """Generate a grounded answer from retrieved Search chunks."""
    contexts = retrieve(question, filter_expression=filter_expression, top=top)
    if not contexts:
        return "I could not find matching indexed sources for this question and access scope."

    context_text = "\n\n".join(
        f"Source: {item['source_file']} chunk {item['chunk_id']}\n{item['content']}"
        for item in contexts
    )
    response = get_openai_client().chat.completions.create(
        model=require_env("MODEL_DEPLOYMENT_NAME"),
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an enterprise RAG assistant. Answer only from the provided context. "
                    "If the context is insufficient, say so. Cite source_file and chunk_id for every factual claim."
                ),
            },
            {"role": "user", "content": f"Question: {question}\n\nContext:\n{context_text}"},
        ],
        temperature=0,
    )
    return response.choices[0].message.content or ""


def evaluate(eval_file: str | None = None) -> list[dict]:
    """Run keyword-based checks from the configured JSONL evaluation file."""
    path = Path(eval_file or os.getenv("RAG_EVAL_FILE", "./eval/eval_questions.jsonl"))
    if not path.exists():
        raise RuntimeError(f"Evaluation file was not found: {path}")

    results = []
    for line_number, line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
        if not line.strip():
            continue
        case = json.loads(line)
        answer = answer_question(case["question"], filter_expression=case.get("filter"))
        missing = [
            keyword
            for keyword in case.get("expected_keywords", [])
            if keyword.lower() not in answer.lower()
        ]
        results.append(
            {
                "line": line_number,
                "question": case["question"],
                "passed": not missing,
                "missing_keywords": missing,
                "answer": answer,
            }
        )

    passed = sum(1 for item in results if item["passed"])
    print(f"Evaluation passed {passed}/{len(results)} cases.")
    for item in results:
        status = "PASS" if item["passed"] else "FAIL"
        print(f"{status}: {item['question']}")
        if item["missing_keywords"]:
            print(f"  Missing: {', '.join(item['missing_keywords'])}")
    return results
