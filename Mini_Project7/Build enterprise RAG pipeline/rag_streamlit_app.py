"""Streamlit enterprise RAG application.

This file is intentionally written as a single runnable app for classroom-demo use.
The main flow is:

1. Load configuration from `.env`.
2. Create cached Azure OpenAI and Azure AI Search clients.
3. Validate that the configured Search index has the fields expected by the app.
4. Route each user question to the right path:
   - enterprise RAG retrieval over Azure AI Search,
   - direct LLM response,
   - live weather lookup,
   - live news lookup,
   - simple sports lookup plus optional news fallback.
5. Render the final answer and supporting source cards in Streamlit.

The enterprise RAG path uses the already-ingested index created by
`Enterprise_RAG_Pipeline_Notebook.ipynb`. For named document requests, the app
can also live-read SharePoint, Blob, or local files as a fallback when the
document is not found in Azure AI Search.
"""

import os
import re
import string
from io import BytesIO
from pathlib import Path
from html import escape
from time import perf_counter

import fitz
import openpyxl
import requests
import streamlit as st
from azure.core.credentials import AzureKeyCredential
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from azure.search.documents import SearchClient
from azure.search.documents.indexes import SearchIndexClient
from azure.storage.blob import ContainerClient
from docx import Document
from dotenv import load_dotenv
from openai import AzureOpenAI


load_dotenv(dotenv_path=".env", override=True)

# These fields are the contract between the ingestion notebook and this app.
# The app checks them at startup so schema mistakes fail visibly before a user
# asks a question.
REQUIRED_FIELDS = {
    "content",
    "content_vector",
    "title",
    "source_file",
    "department",
    "security_level",
    "chunk_id",
}

SUPPORTED_LIVE_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".log",
    ".pdf",
    ".docx",
    ".xlsx",
    ".xlsm",
}

# Quick prompts shown as sidebar buttons. They exercise the major routes:
# HR RAG, IT RAG, mixed enterprise RAG, direct LLM, weather, and news.
SAMPLE_QUESTIONS = [
    "Does the ABC Corp dataset mention employee records?",
    "Summarize the HR leave policy.",
    "What IT security rules are mentioned?",
    "What benefits are available to employees?",
    "Compare HR and IT policy responsibilities.",
    "Draft a short email asking employees to follow company policy.",
    "What is the weather in Bangalore?",
    "Show latest AI news.",
]

GREETING_MESSAGES = {
    "hi",
    "hello",
    "hey",
    "hi there",
    "hello there",
    "hey there",
    "how are you",
    "hi how are you",
}

# Routing is deliberately keyword-based for a transparent training demo. In a
# production app this can be replaced by a classifier, policy engine, or agent.
HR_KEYWORDS = {
    "hr",
    "human resource",
    "human resources",
    "employee",
    "employees",
    "payroll",
    "salary",
    "leave",
    "benefit",
    "benefits",
    "onboarding",
    "offboarding",
    "attendance",
    "manager",
    "designation",
    "joining",
    "records",
}

IT_KEYWORDS = {
    "helpdesk",
    "help desk",
    "support",
    "ticket",
    "password",
    "laptop",
    "vpn",
    "network",
    "security",
    "system",
    "software",
    "hardware",
    "access",
    "incident",
    "email issue",
    "phishing",
}

DOCUMENT_LOOKUP_KEYWORDS = {
    "azure blob",
    "blob",
    "sharepoint",
    "document",
    "documents",
    "doc",
    "docs",
    "file",
    "pdf",
    "guide",
    "summarize",
    "summary",
}

MIXED_KEYWORDS = {
    "all",
    "both",
    "compare",
    "combined",
    "company",
    "enterprise",
    "organization",
    "policies",
    "policy summary",
    "overall",
}

WEATHER_KEYWORDS = {
    "weather",
    "temperature",
    "forecast",
    "rain",
    "humidity",
    "wind",
}

NEWS_KEYWORDS = {
    "news",
    "headlines",
    "latest",
    "breaking",
    "updates",
}

SPORTS_KEYWORDS = {
    "ipl",
    "cricket",
    "t20",
    "match",
    "final",
    "champion",
    "champions",
    "rcb",
    "gt",
    "gujarat titans",
    "royal challengers",
}

NEWS_PROVIDERS = [
    {
        "name": "NewsAPI",
        "env_key": "NEWSAPI_KEY",
        "url": "https://newsapi.org/v2/everything",
        "params": lambda topic, key: {
            "q": topic,
            "language": "en",
            "sortBy": "publishedAt",
            "pageSize": 5,
            "apiKey": key,
        },
        "parser": "newsapi",
    },
    {
        "name": "GNews",
        "env_key": "GNEWS_API_KEY",
        "url": "https://gnews.io/api/v4/search",
        "params": lambda topic, key: {
            "q": topic,
            "lang": "en",
            "max": 5,
            "apikey": key,
        },
        "parser": "gnews",
    },
    {
        "name": "TheNewsAPI",
        "env_key": "THENEWSAPI_KEY",
        "url": "https://api.thenewsapi.com/v1/news/all",
        "params": lambda topic, key: {
            "search": topic,
            "language": "en",
            "sort": "published_on",
            "limit": 5,
            "api_token": key,
        },
        "parser": "thenewsapi",
    },
]

# Lightweight static lookup used for simple sports questions. Current-event
# sports questions can still fall through to the configured news provider.
IPL_WINNERS = {
    "2008": "Rajasthan Royals",
    "2009": "Deccan Chargers",
    "2010": "Chennai Super Kings",
    "2011": "Chennai Super Kings",
    "2012": "Kolkata Knight Riders",
    "2013": "Mumbai Indians",
    "2014": "Kolkata Knight Riders",
    "2015": "Mumbai Indians",
    "2016": "Sunrisers Hyderabad",
    "2017": "Mumbai Indians",
    "2018": "Chennai Super Kings",
    "2019": "Mumbai Indians",
    "2020": "Mumbai Indians",
    "2021": "Chennai Super Kings",
    "2022": "Gujarat Titans",
    "2023": "Chennai Super Kings",
    "2024": "Kolkata Knight Riders",
    "2025": "Royal Challengers Bengaluru",
    "2026": "Royal Challengers Bengaluru",
}


def require_env(name: str) -> str:
    """Return a required `.env` value or raise a readable app error.

    Many Azure SDK failures become noisy if configuration is missing. This
    helper keeps startup and request-time errors focused on the exact missing
    environment variable.
    """
    value = os.getenv(name)
    if not value or value.startswith("<"):
        raise RuntimeError(f"Missing required .env value: {name}")
    return value


def optional_env(name: str) -> str | None:
    """Return a non-placeholder `.env` value when present."""
    value = os.getenv(name)
    if not value or value.startswith("<"):
        return None
    return value


def live_fallback_enabled() -> bool:
    """Return whether live source fallback is enabled.

    The app defaults to enabled because this demo now supports live reads. Set
    `RAG_LIVE_FALLBACK=false` to force indexed-only behavior.
    """
    return os.getenv("RAG_LIVE_FALLBACK", "true").strip().lower() not in {"0", "false", "no", "off"}


def configured_rag_sources() -> list[str]:
    """Read `RAG_SOURCE` and return normalized source names.

    `all` expands to SharePoint, Blob, and local. Unknown source values are
    ignored so a typo does not crash unrelated app routes.
    """
    raw_value = os.getenv("RAG_SOURCE", "local")
    sources = [item.strip().lower() for item in raw_value.split(",") if item.strip()]
    if "all" in sources:
        return ["sharepoint", "blob", "local"]
    return [source for source in sources if source in {"sharepoint", "blob", "local"}]


def preferred_live_sources(question: str) -> list[str]:
    """Return live sources in the order that best matches the user's wording.

    `RAG_SOURCE` defines the default order. If the user explicitly says Blob,
    SharePoint, or local, that source is checked first for this request. This
    matters when the same filename exists in multiple live sources.
    """
    sources = configured_rag_sources()
    normalized = normalize_query(question)
    preferences = []
    if "blob" in normalized or "azure blob" in normalized:
        preferences.append("blob")
    if "sharepoint" in normalized:
        preferences.append("sharepoint")
    if "local" in normalized or "folder" in normalized:
        preferences.append("local")

    ordered = [source for source in preferences if source in sources]
    ordered.extend(source for source in sources if source not in ordered)
    return ordered


def is_greeting(text: str) -> bool:
    """Detect simple greetings so the app can avoid unnecessary RAG calls."""
    normalized = normalize_query(text)
    return normalized in GREETING_MESSAGES


def normalize_query(text: str) -> str:
    """Normalize user text for keyword routing.

    The function lowercases text, separates letters from digits, removes
    punctuation, and collapses whitespace. This makes rules like "IPL2025" and
    "IPL 2025?" behave the same way.
    """
    lowered = text.lower()
    lowered = re.sub(r"([a-z])(\d)", r"\1 \2", lowered)
    lowered = re.sub(r"(\d)([a-z])", r"\1 \2", lowered)
    return " ".join(lowered.translate(str.maketrans("", "", string.punctuation)).split())


def normalize_document_name(text: str) -> str:
    """Normalize file/title text for document-name matching.

    Azure AI Search may not rank a filename-style request highly when the user
    includes greeting words or extra instructions. This helper lets the app
    compare names such as `Professional_Etiquette_in_IT` and
    `Professional Etiquette in IT.pdf` using the same representation.
    """
    return normalize_query(text.replace("_", " ").replace("-", " "))


def extract_requested_document_hint(question: str) -> str | None:
    """Extract an explicit filename-like document hint from the question.

    The app only treats clear filename-style tokens as hard document hints. This
    avoids blocking normal broad questions like "summarize HR policy documents"
    while still handling prompts such as `Professional_Etiquette_in_IT`.
    """
    quoted = re.findall(r"['\"]([^'\"]{4,})['\"]", question)
    if quoted:
        return max(quoted, key=len)

    filename_tokens = re.findall(r"\b[\w.-]*[_][\w.-]*\b", question)
    if filename_tokens:
        return max(filename_tokens, key=len)

    marker_match = re.search(
        r"(?:sharepoint|blob|local|file|document)\s*:?\s+(.+?)(?:\s+(?:document|file|pdf|docx|xlsx|xlsm|summari[sz]e|details?|please)\b|$)",
        question,
        flags=re.IGNORECASE,
    )
    if marker_match:
        candidate = marker_match.group(1).strip(" :;,.\"'")
        if len(candidate.split()) >= 2:
            return candidate

    return None


def contains_keyword(normalized: str, keywords: set[str]) -> bool:
    """Check whether any full keyword or phrase exists in normalized text."""
    return any(re.search(rf"(?<!\w){re.escape(keyword)}(?!\w)", normalized) for keyword in keywords)


def odata_quote(value: str) -> str:
    """Escape single quotes for Azure AI Search OData filter strings."""
    return value.replace("'", "''")


def build_access_filter(department: str, security_level: str) -> str:
    """Build the server-side Azure AI Search filter.

    `security_level` is always enforced. `department` is enforced unless the UI
    or router selects `All`. In production, these values should come from
    signed-in user claims or groups, not free-form UI choices.
    """
    filters = [f"security_level eq '{odata_quote(security_level)}'"]
    if department != "All":
        filters.append(f"department eq '{odata_quote(department)}'")
    return " and ".join(filters)


def route_query(question: str) -> dict:
    """Choose the execution route for a user question.

    Return shape:
    - `route`: user-visible route label.
    - `department`: Search filter department for enterprise RAG routes.
    - `retrieval`: whether Azure AI Search retrieval is required.
    - `tool`: optional external tool path such as weather or news.
    - `reason`: short explanation shown in the UI.
    """
    normalized = normalize_query(question)
    has_weather = contains_keyword(normalized, WEATHER_KEYWORDS)
    has_news = contains_keyword(normalized, NEWS_KEYWORDS)
    has_sports = contains_keyword(normalized, SPORTS_KEYWORDS)
    has_hr = contains_keyword(normalized, HR_KEYWORDS)
    has_it = contains_keyword(normalized, IT_KEYWORDS)
    has_document_lookup = contains_keyword(normalized, DOCUMENT_LOOKUP_KEYWORDS)
    has_mixed = contains_keyword(normalized, MIXED_KEYWORDS)

    if has_weather:
        return {
            "route": "weather",
            "department": "All",
            "retrieval": False,
            "tool": "weather",
            "reason": "Weather question. Calling the live weather API.",
        }
    if has_news:
        return {
            "route": "news",
            "department": "All",
            "retrieval": False,
            "tool": "news",
            "reason": "News question. Calling the configured news API.",
        }
    if has_sports:
        return {
            "route": "sports-news",
            "department": "All",
            "retrieval": False,
            "tool": "sports",
            "reason": "Sports/current-events question. Checking sports lookup and configured news API.",
        }
    if has_mixed or (has_hr and has_it):
        return {
            "route": "mixed",
            "department": "All",
            "retrieval": True,
            "tool": None,
            "reason": "Mixed or enterprise-wide question. Searching all permitted departments.",
        }
    if has_document_lookup and not (has_hr or has_it):
        return {
            "route": "document-lookup",
            "department": "All",
            "retrieval": True,
            "tool": None,
            "reason": "Document lookup question. Searching all indexed sources.",
        }
    if has_hr:
        return {
            "route": "hr",
            "department": "HR",
            "retrieval": True,
            "tool": None,
            "reason": "HR-related question. Searching HR documents only.",
        }
    if has_it:
        return {
            "route": "it",
            "department": "IT",
            "retrieval": True,
            "tool": None,
            "reason": "IT/helpdesk-related question. Searching IT documents only.",
        }
    return {
        "route": "generic",
        "department": "All",
        "retrieval": False,
        "tool": None,
        "reason": "General question. Answering directly with the LLM without document retrieval.",
    }


@st.cache_resource(show_spinner=False)
def get_openai_client() -> AzureOpenAI:
    """Create one cached Azure OpenAI client for the Streamlit process.

    Streamlit reruns the script on every interaction. `st.cache_resource`
    prevents the app from recreating network clients every time the user types
    a message or clicks a button.

    Authentication uses Microsoft Entra ID through `DefaultAzureCredential`.
    The signed-in identity must have access to the Azure OpenAI resource.
    """
    token_provider = get_bearer_token_provider(
        DefaultAzureCredential(),
        "https://cognitiveservices.azure.com/.default",
    )
    return AzureOpenAI(
        azure_endpoint=require_env("AZURE_OPENAI_ENDPOINT"),
        azure_ad_token_provider=token_provider,
        api_version="2024-10-21",
    )


@st.cache_resource(show_spinner=False)
def get_search_client() -> SearchClient:
    """Create one cached Azure AI Search document client.

    `SearchClient` is used for runtime document queries against one configured
    index. This app uses an API key because it is simple for demos; production
    deployments should prefer managed identity or another Entra-based pattern.
    """
    return SearchClient(
        endpoint=require_env("AI_SEARCH_SERVICE_ENDPOINT"),
        index_name=require_env("AI_SEARCH_INDEX_NAME"),
        credential=AzureKeyCredential(require_env("AI_SEARCH_API_KEY")),
    )


@st.cache_resource(show_spinner=False)
def get_index_client() -> SearchIndexClient:
    """Create one cached Azure AI Search index-management client.

    The app uses this client only to inspect the index schema at startup. It
    does not create, update, or delete indexes.
    """
    return SearchIndexClient(
        endpoint=require_env("AI_SEARCH_SERVICE_ENDPOINT"),
        credential=AzureKeyCredential(require_env("AI_SEARCH_API_KEY")),
    )


@st.cache_data(ttl=120, show_spinner=False)
def get_index_summary() -> dict:
    """Read index health information for the dashboard cards.

    This verifies the schema contract and samples up to 100 indexed chunks so
    the UI can show source files, departments, and a quick document count. The
    short cache keeps the UI responsive while still refreshing during demos.
    """
    index_name = require_env("AI_SEARCH_INDEX_NAME")
    index = get_index_client().get_index(index_name)
    existing_fields = {field.name for field in index.fields}
    missing_fields = sorted(REQUIRED_FIELDS - existing_fields)
    docs = list(
        get_search_client().search(
            search_text="*",
            select=["source_file", "department", "security_level", "chunk_id"],
            top=100,
        )
    )
    sources = sorted({doc["source_file"] for doc in docs})
    departments = sorted({doc["department"] for doc in docs})
    return {
        "index_name": index_name,
        "document_count": len(docs),
        "source_count": len(sources),
        "sources": sources,
        "departments": departments,
        "schema_ok": not missing_fields,
        "missing_fields": missing_fields,
    }


def embed_question(question: str) -> list[float]:
    """Generate the query embedding used for vector retrieval.

    The embedding deployment must produce the same vector dimension as the
    Search index field `content_vector`; in this project that dimension is 1536.
    """
    response = get_openai_client().embeddings.create(
        model=require_env("EMBEDDING_DEPLOYMENT_NAME"),
        input=[question],
    )
    return response.data[0].embedding


def retrieve(question: str, filter_expression: str, top: int) -> list[dict]:
    """Retrieve the best matching chunks from Azure AI Search.

    The query combines:
    - `search_text=question` for keyword/semantic text matching,
    - `vector_queries` over `content_vector` for embedding similarity,
    - `query_type="semantic"` for semantic ranking,
    - `filter=filter_expression` for access control and department scoping.

    This is app-side vector retrieval. It is different from the Foundry Agent
    notebook's integrated-vectorizer requirement because here the app generates
    the query embedding itself before calling Azure AI Search.
    """
    question_vector = embed_question(question)
    results = get_search_client().search(
        search_text=question,
        vector_queries=[
            {
                "kind": "vector",
                "vector": question_vector,
                "fields": "content_vector",
                "k": top,
            }
        ],
        query_type="semantic",
        semantic_configuration_name="rag-semantic-config",
        filter=filter_expression,
        top=top,
        select=["content", "title", "source_file", "department", "security_level", "chunk_id"],
    )
    return [dict(item) for item in results]


def retrieve_named_document(question: str, filter_expression: str, top: int) -> tuple[list[dict], str | None, list[str]]:
    """Try to retrieve chunks from a specifically named document first.

    Returns `(contexts, requested_hint, available_sources)`.

    This is useful when the user asks for a document by filename/title. The app
    scans source metadata under the active access filter and matches the
    requested name against `source_file` and `title`. If it finds matches, those
    chunks are used as context instead of relying on general semantic ranking.
    """
    requested_hint = extract_requested_document_hint(question)
    if not requested_hint:
        return [], None, []

    normalized_hint = normalize_document_name(requested_hint)
    results = get_search_client().search(
        search_text="*",
        filter=filter_expression,
        top=1000,
        select=["content", "title", "source_file", "department", "security_level", "chunk_id"],
    )
    docs = [dict(item) for item in results]
    available_sources = sorted({str(item.get("source_file") or "Unknown source") for item in docs})

    matches = []
    for item in docs:
        source_name = normalize_document_name(str(item.get("source_file") or ""))
        title = normalize_document_name(str(item.get("title") or ""))
        if normalized_hint in source_name or normalized_hint in title:
            matches.append(item)

    matches.sort(key=lambda item: int(item.get("chunk_id") or 0))
    return matches[:top], requested_hint, available_sources


def parse_live_document(file_name: str, content: bytes) -> str:
    """Extract text from a live SharePoint, Blob, or local document.

    This parser mirrors the ingestion notebook's common document types. It runs
    in memory and does not update Azure AI Search.
    """
    suffix = Path(file_name).suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".log"}:
        return content.decode("utf-8", errors="ignore")

    if suffix == ".pdf":
        with fitz.open(stream=content, filetype="pdf") as pdf:
            return "\n\n".join(page.get_text("text") for page in pdf)

    if suffix == ".docx":
        document = Document(BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)

    if suffix in {".xlsx", ".xlsm"}:
        workbook = openpyxl.load_workbook(BytesIO(content), data_only=True, read_only=True)
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value) for value in row if value is not None]
                if values:
                    lines.append(" | ".join(values))
        workbook.close()
        return "\n".join(lines)

    raise ValueError(f"Live read does not support `{suffix}` files yet.")


def build_live_context(file_name: str, content: bytes, source_label: str, max_chars: int = 16000) -> dict:
    """Parse a live document and return one source-card compatible context."""
    text = parse_live_document(file_name, content).strip()
    if not text:
        raise ValueError(f"`{file_name}` was found in {source_label}, but no text could be extracted.")

    return {
        "content": text[:max_chars],
        "title": Path(file_name).stem,
        "source_file": f"{source_label}: {file_name}",
        "department": "Live",
        "security_level": "source-controlled",
        "chunk_id": "live",
    }


@st.cache_data(ttl=180, show_spinner=False)
def list_local_live_files() -> list[dict]:
    """List supported local files from `RAG_DOCS_DIR` for live fallback."""
    docs_dir = Path(os.getenv("RAG_DOCS_DIR", "./data/documents"))
    if not docs_dir.exists():
        return []

    return [
        {"name": path.name, "path": str(path)}
        for path in docs_dir.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_LIVE_EXTENSIONS
    ]


def read_local_live_file(file_info: dict) -> bytes:
    """Read a matched local live file."""
    return Path(file_info["path"]).read_bytes()


def get_blob_container_client() -> ContainerClient:
    """Create a Blob container client from `BLOB_CONTAINER_URL`.

    SAS URLs authenticate themselves. Non-SAS URLs use `DefaultAzureCredential`,
    so the signed-in identity needs Blob data reader access.
    """
    container_url = require_env("BLOB_CONTAINER_URL")
    if "sig=" in container_url.lower():
        return ContainerClient.from_container_url(container_url)
    return ContainerClient.from_container_url(container_url, credential=DefaultAzureCredential())


@st.cache_data(ttl=180, show_spinner=False)
def list_blob_live_files() -> list[dict]:
    """List supported blobs from the configured container for live fallback."""
    container = get_blob_container_client()
    files = []
    for blob in container.list_blobs():
        name = blob.name
        if Path(name).suffix.lower() in SUPPORTED_LIVE_EXTENSIONS:
            files.append({"name": name, "path": name})
    return files


def read_blob_live_file(file_info: dict) -> bytes:
    """Download a matched blob into memory."""
    container = get_blob_container_client()
    return container.download_blob(file_info["path"]).readall()


def get_sharepoint_token() -> str:
    """Acquire a Microsoft Graph token using the SharePoint app registration."""
    tenant_id = require_env("AZURE_TENANT_ID")
    response = requests.post(
        f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token",
        data={
            "client_id": require_env("SHAREPOINT_CLIENT_ID"),
            "client_secret": require_env("SHAREPOINT_CLIENT_SECRET"),
            "scope": "https://graph.microsoft.com/.default",
            "grant_type": "client_credentials",
        },
        timeout=30,
    )
    response.raise_for_status()
    return response.json()["access_token"]


def graph_get(url: str, token: str) -> dict:
    """Call Microsoft Graph and return JSON."""
    response = requests.get(url, headers={"Authorization": f"Bearer {token}"}, timeout=30)
    response.raise_for_status()
    return response.json()


@st.cache_data(ttl=180, show_spinner=False)
def list_sharepoint_live_files() -> list[dict]:
    """List supported files from the configured SharePoint document library."""
    token = get_sharepoint_token()
    hostname = require_env("SHAREPOINT_HOSTNAME")
    site_path = require_env("SHAREPOINT_SITE_PATH")
    drive_name = os.getenv("SHAREPOINT_DRIVE_NAME", "Documents")

    site = graph_get(f"https://graph.microsoft.com/v1.0/sites/{hostname}:{site_path}", token)
    drives = graph_get(f"https://graph.microsoft.com/v1.0/sites/{site['id']}/drives", token).get("value", [])
    drive = next((item for item in drives if item.get("name") == drive_name), None)
    if not drive:
        raise RuntimeError(f"SharePoint drive `{drive_name}` was not found.")

    files = []
    stack = [f"https://graph.microsoft.com/v1.0/drives/{drive['id']}/root/children"]
    while stack:
        payload = graph_get(stack.pop(), token)
        for item in payload.get("value", []):
            if "folder" in item:
                stack.append(f"https://graph.microsoft.com/v1.0/drives/{drive['id']}/items/{item['id']}/children")
            elif "file" in item and Path(item.get("name", "")).suffix.lower() in SUPPORTED_LIVE_EXTENSIONS:
                files.append(
                    {
                        "name": item["name"],
                        "path": item.get("parentReference", {}).get("path", ""),
                        "id": item["id"],
                        "drive_id": drive["id"],
                    }
                )
        next_link = payload.get("@odata.nextLink")
        if next_link:
            stack.append(next_link)
    return files


def read_sharepoint_live_file(file_info: dict) -> bytes:
    """Download a matched SharePoint file through Microsoft Graph."""
    token = get_sharepoint_token()
    url = f"https://graph.microsoft.com/v1.0/drives/{file_info['drive_id']}/items/{file_info['id']}/content"
    response = requests.get(url, headers={"Authorization": f"Bearer {token}"}, timeout=60)
    response.raise_for_status()
    return response.content


def match_live_file(files: list[dict], requested_hint: str) -> dict | None:
    """Find the best live file match for the requested document name."""
    normalized_hint = normalize_document_name(requested_hint)
    exact_matches = []
    partial_matches = []
    for file_info in files:
        normalized_name = normalize_document_name(file_info["name"])
        normalized_stem = normalize_document_name(Path(file_info["name"]).stem)
        if normalized_hint == normalized_stem or normalized_hint == normalized_name:
            exact_matches.append(file_info)
        elif normalized_hint in normalized_stem or normalized_hint in normalized_name:
            partial_matches.append(file_info)
    return (exact_matches or partial_matches or [None])[0]


def retrieve_live_named_document(question: str) -> tuple[list[dict], str | None, list[str], list[str]]:
    """Find and parse a named document directly from configured live sources.

    Returns `(contexts, requested_hint, checked_sources, errors)`.
    """
    requested_hint = extract_requested_document_hint(question)
    if not requested_hint or not live_fallback_enabled():
        return [], requested_hint, [], []

    checked_sources = []
    errors = []
    source_loaders = {
        "local": (list_local_live_files, read_local_live_file, "Local file"),
        "blob": (list_blob_live_files, read_blob_live_file, "Blob live"),
        "sharepoint": (list_sharepoint_live_files, read_sharepoint_live_file, "SharePoint live"),
    }

    for source_name in preferred_live_sources(question):
        list_files, read_file, source_label = source_loaders[source_name]
        checked_sources.append(source_label)
        try:
            files = list_files()
            match = match_live_file(files, requested_hint)
            if not match:
                continue
            content = read_file(match)
            return [build_live_context(match["name"], content, source_label)], requested_hint, checked_sources, errors
        except Exception as exc:
            errors.append(f"{source_label}: {exc}")

    return [], requested_hint, checked_sources, errors


def answer_question(question: str, filter_expression: str, top: int) -> dict:
    """Run the full enterprise RAG answer path.

    Steps:
    1. Retrieve relevant chunks using the access filter.
    2. Pack the retrieved chunks into a compact context string.
    3. Ask the chat model to answer only from that context.
    4. Return both the answer and raw source metadata for UI source cards.
    """
    named_contexts, requested_hint, available_sources = retrieve_named_document(question, filter_expression, top)
    if named_contexts:
        contexts = named_contexts
    elif requested_hint:
        live_contexts, _, checked_sources, live_errors = retrieve_live_named_document(question)
        if live_contexts:
            contexts = live_contexts
        else:
            checked_text = ", ".join(checked_sources) if checked_sources else "none"
            error_text = "\n".join(f"- {error}" for error in live_errors)
            if error_text:
                error_text = f"\n\nLive source errors:\n{error_text}"
            source_list = "\n".join(f"- {source}" for source in available_sources[:25]) or "- No sources visible for this access scope."
            return {
                "answer": (
                    f"I could not find `{requested_hint}` in the current Azure AI Search index or configured live sources. "
                    f"Checked live sources: {checked_text}.\n\n"
                    "If the document was just added, confirm `RAG_SOURCE` includes its source and that the app has permission to read it.\n\n"
                    f"Indexed sources currently visible:\n{source_list}"
                    f"{error_text}"
                ),
                "sources": [],
            }
    else:
        contexts = retrieve(question, filter_expression, top)

    if requested_hint and contexts and any(str(item.get("chunk_id")) == "live" for item in contexts):
        live_note = (
            "This answer uses live source content that was read directly from SharePoint, Blob, or local storage. "
            "It has not necessarily been uploaded to Azure AI Search yet.\n\n"
        )
    else:
        live_note = ""

    if not contexts:
        return {
            "answer": (
                "I could not find matching indexed sources for this question and access scope. "
                "Try another department, choose All, or confirm the documents were ingested."
            ),
            "sources": [],
        }

    context_text = "\n\n".join(
        f"Source: {item['source_file']} chunk {item['chunk_id']}\n{item['content']}"
        for item in contexts
    )
    response = get_openai_client().chat.completions.create(
        model=require_env("MODEL_DEPLOYMENT_NAME"),
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an enterprise RAG assistant for daily employee use. "
                    "Answer only from the provided context. If the context is insufficient, say so. "
                    "Cite source_file and chunk_id for every factual claim. Keep answers concise and business-ready."
                ),
            },
            {"role": "user", "content": f"Question: {question}\n\nContext:\n{context_text}"},
        ],
        temperature=0,
    )
    return {"answer": live_note + response.choices[0].message.content, "sources": contexts}


def answer_generic_question(question: str) -> dict:
    """Answer non-enterprise questions directly with the chat model.

    This path intentionally does not claim document grounding because no search
    context is supplied to the model.
    """
    response = get_openai_client().chat.completions.create(
        model=require_env("MODEL_DEPLOYMENT_NAME"),
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a concise enterprise assistant. Answer general questions directly. "
                    "Do not pretend to have checked company documents unless retrieval context is provided."
                ),
            },
            {"role": "user", "content": question},
        ],
        temperature=0.2,
    )
    return {"answer": response.choices[0].message.content, "sources": []}


def extract_location(question: str) -> str:
    """Extract a simple location phrase from a weather question."""
    normalized = question.strip().strip("?")
    lowered = normalized.lower()
    for marker in [" in ", " at ", " for "]:
        if marker in lowered:
            value = normalized[lowered.rfind(marker) + len(marker) :].strip()
            return value or "Bangalore"
    return "Bangalore"


def answer_weather_question(question: str) -> dict:
    """Answer weather questions with the live Open-Meteo API.

    Open-Meteo requires latitude and longitude, so the function first calls the
    geocoding endpoint and then calls the forecast endpoint for current weather.
    """
    location = extract_location(question)
    country_code = None
    if location.lower() in {"bangalore", "bengaluru", "bangalore india", "bengaluru india"}:
        location = "Bengaluru"
        country_code = "IN"
    elif " india" in location.lower():
        location = location.lower().replace(" india", "").strip().title()
        country_code = "IN"

    geo_params = {"name": location, "count": 1, "language": "en", "format": "json"}
    if country_code:
        geo_params["countryCode"] = country_code

    geo_response = requests.get(
        "https://geocoding-api.open-meteo.com/v1/search",
        params=geo_params,
        timeout=15,
    )
    geo_response.raise_for_status()
    matches = geo_response.json().get("results") or []
    if not matches:
        return {"answer": f"I could not find a weather location match for '{location}'.", "sources": []}

    place = matches[0]
    weather_response = requests.get(
        "https://api.open-meteo.com/v1/forecast",
        params={
            "latitude": place["latitude"],
            "longitude": place["longitude"],
            "current": "temperature_2m,relative_humidity_2m,apparent_temperature,precipitation,wind_speed_10m",
            "timezone": "auto",
        },
        timeout=15,
    )
    weather_response.raise_for_status()
    payload = weather_response.json()
    current = payload.get("current", {})
    units = payload.get("current_units", {})
    place_name = ", ".join(
        value
        for value in [place.get("name"), place.get("admin1"), place.get("country")]
        if value
    )

    answer = (
        f"Current weather for {place_name}:\n\n"
        f"- Temperature: {current.get('temperature_2m')} {units.get('temperature_2m', '')}\n"
        f"- Feels like: {current.get('apparent_temperature')} {units.get('apparent_temperature', '')}\n"
        f"- Humidity: {current.get('relative_humidity_2m')} {units.get('relative_humidity_2m', '')}\n"
        f"- Precipitation: {current.get('precipitation')} {units.get('precipitation', '')}\n"
        f"- Wind speed: {current.get('wind_speed_10m')} {units.get('wind_speed_10m', '')}\n"
        f"- Observation time: {current.get('time')}\n\n"
        "Source: Open-Meteo live weather API."
    )
    return {
        "answer": answer,
        "sources": [
            {
                "source_file": "Open-Meteo Weather API",
                "chunk_id": "live",
                "department": "External",
                "security_level": "public",
                "content": f"Weather API result for {place_name}: {current}",
            }
        ],
    }


def extract_news_topic(question: str) -> str:
    """Reduce a news-style question to a provider search topic."""
    cleaned = normalize_query(question)
    if "ipl" in cleaned:
        year = extract_year(question)
        if year:
            return f"IPL {year} final winner"
        return "IPL final winner"

    stop_words = {"latest", "news", "headlines", "breaking", "updates", "show", "me", "which", "team"}
    topic = " ".join(word for word in cleaned.split() if word not in stop_words)
    return topic or "technology"


def extract_year(question: str) -> str | None:
    """Extract a four-digit year from normalized user text."""
    match = re.search(r"\b(20\d{2})\b", normalize_query(question))
    return match.group(1) if match else None


def get_configured_news_provider() -> dict | None:
    """Return the first configured news provider, if any.

    Kept as a convenience helper even though the answer path uses all configured
    providers for fallback.
    """
    providers = get_configured_news_providers()
    return providers[0] if providers else None


def get_configured_news_providers() -> list[dict]:
    """Return news providers that have API keys in `.env`.

    If `NEWS_PROVIDER` is set, only that provider is considered. If it is not
    set, the app tries all providers with configured keys in the order defined
    by `NEWS_PROVIDERS`.
    """
    requested_provider = os.getenv("NEWS_PROVIDER", "").strip().lower()
    providers = NEWS_PROVIDERS
    if requested_provider:
        providers = [provider for provider in NEWS_PROVIDERS if provider["name"].lower() == requested_provider]

    configured = []
    for provider in providers:
        api_key = os.getenv(provider["env_key"])
        if api_key:
            configured.append({**provider, "api_key": api_key})
    return configured


def news_timeout_seconds() -> int:
    """Read and clamp the news API timeout setting."""
    raw_value = os.getenv("NEWS_API_TIMEOUT_SECONDS", "45").strip()
    try:
        return max(10, min(int(raw_value), 90))
    except ValueError:
        return 45


def normalize_news_articles(provider: dict, payload: dict) -> list[dict]:
    """Convert each provider's response shape into one common article list."""
    parser = provider["parser"]
    if parser in {"newsapi", "gnews"}:
        return payload.get("articles", [])[:5]
    if parser == "thenewsapi":
        articles = payload.get("data", [])[:5]
        return [
            {
                "title": article.get("title"),
                "description": article.get("description") or article.get("snippet"),
                "publishedAt": article.get("published_at"),
                "url": article.get("url"),
                "source": {"name": article.get("source")},
            }
            for article in articles
        ]
    return []


def answer_news_question(question: str) -> dict:
    """Answer live news questions with configured free-tier news APIs.

    The function tries each configured provider until one succeeds. Source cards
    contain article summaries and source names rather than internal documents.
    """
    providers = get_configured_news_providers()
    if not providers:
        return {
            "answer": (
                "News API is not configured yet. Add one of these free-tier API keys in `.env` to enable live news/sports lookup:\n\n"
                "- `NEWSAPI_KEY` for NewsAPI.org, 100 requests/day free tier\n"
                "- `GNEWS_API_KEY` for GNews API, 100 requests/day free tier\n"
                "- `THENEWSAPI_KEY` for TheNewsAPI free tier\n\n"
                "Optional: set `NEWS_PROVIDER=\"NewsAPI\"`, `NEWS_PROVIDER=\"GNews\"`, or `NEWS_PROVIDER=\"TheNewsAPI\"`."
            ),
            "sources": [],
        }

    topic = extract_news_topic(question)
    errors = []
    for provider in providers:
        try:
            response = requests.get(
                provider["url"],
                params=provider["params"](topic, provider["api_key"]),
                timeout=(8, news_timeout_seconds()),
            )
            response.raise_for_status()
            articles = normalize_news_articles(provider, response.json())
        except requests.exceptions.Timeout:
            errors.append(f"{provider['name']} timed out")
            continue
        except requests.exceptions.RequestException as exc:
            errors.append(f"{provider['name']} error: {exc}")
            continue

        if not articles:
            return {"answer": f"No recent news articles found for '{topic}' from {provider['name']}.", "sources": []}

        lines = [f"Latest results for '{topic}' from {provider['name']}:"]
        sources = []
        for index, article in enumerate(articles, start=1):
            title = article.get("title") or "Untitled"
            source_name = (article.get("source") or {}).get("name") or "Unknown source"
            published = article.get("publishedAt") or "unknown date"
            url = article.get("url") or ""
            lines.append(f"{index}. {title} ({source_name}, {published})\n   {url}")
            sources.append(
                {
                    "source_file": source_name,
                    "chunk_id": index,
                    "department": "External",
                    "security_level": "public",
                    "content": article.get("description") or title,
                }
            )

        return {"answer": "\n\n".join(lines), "sources": sources}

    return {
        "answer": (
            f"I tried the configured news provider for '{topic}', but it did not respond in time. "
            "Please try again in a minute, or add another provider key such as `NEWSAPI_KEY` or `GNEWS_API_KEY` for fallback."
        ),
        "sources": [
            {
                "source_file": "News tool status",
                "chunk_id": 1,
                "department": "External",
                "security_level": "public",
                "content": "; ".join(errors) if errors else "No configured news provider returned articles.",
            }
        ],
    }


def answer_sports_question(question: str) -> dict:
    """Answer sports questions.

    Simple IPL winner questions are handled from the static lookup for speed.
    Broader or current-event sports questions fall back to live news search.
    """
    normalized = normalize_query(question)
    year = extract_year(question)
    is_ipl_question = contains_keyword(normalized, {"ipl"})
    is_title_question = contains_keyword(normalized, {"winner", "won", "title", "champion", "champions"})

    if is_ipl_question and year in IPL_WINNERS and is_title_question:
        winner = IPL_WINNERS[year]
        answer = f"{winner} won the IPL {year} title."
        if year == "2023":
            answer += " Chennai Super Kings beat Gujarat Titans in the final."
        elif year == "2026":
            answer += " Royal Challengers Bengaluru beat Gujarat Titans by 5 wickets in the final."

        sources = [
            {
                "source_file": "IPL winners lookup",
                "chunk_id": year,
                "department": "External",
                "security_level": "public",
                "content": answer,
            }
        ]
        try:
            news_result = answer_news_question(question)
            sources.extend(news_result.get("sources", [])[:2])
        except Exception:
            pass
        return {"answer": answer, "sources": sources}

    return answer_news_question(question)


def add_css() -> None:
    """Inject Streamlit CSS for the demo UI.

    The app keeps CSS in Python so the demo remains a single runnable file.
    The styles make the page denser and more application-like than Streamlit's
    default layout while preserving the native widgets.
    """
    st.markdown(
        """
        <style>
          :root {
            --ink: #111827;
            --muted: #475569;
            --line: #cbd5e1;
            --panel: #ffffff;
            --page: #f3f6f8;
            --accent: #0f766e;
            --accent-dark: #115e59;
          }

          .stApp {
            background: var(--page) !important;
            color: var(--ink) !important;
            font-family: "Segoe UI", Arial, Helvetica, sans-serif !important;
          }

          [data-testid="stHeader"] {
            background: var(--page) !important;
          }

          .main .block-container {
            max-width: 1240px;
            padding-top: 1.25rem;
            padding-bottom: 3rem;
          }

          [data-testid="stSidebar"] {
            background: #ffffff !important;
            border-right: 1px solid var(--line);
          }

          [data-testid="stSidebar"] h1,
          [data-testid="stSidebar"] h2,
          [data-testid="stSidebar"] h3,
          [data-testid="stSidebar"] p,
          [data-testid="stSidebar"] span,
          [data-testid="stSidebar"] label,
          [data-testid="stSidebar"] div {
            color: var(--ink) !important;
          }

          .hero {
            background: linear-gradient(135deg, #0f766e 0%, #155e63 55%, #25614c 100%);
            border: 1px solid #0d5f59;
            border-radius: 8px;
            padding: 1.25rem 1.4rem;
            margin-bottom: 1rem;
            box-shadow: 0 8px 24px rgba(15, 23, 42, .14);
          }

          .hero-title {
            color: #ffffff !important;
            font-size: 2rem;
            line-height: 1.15;
            font-weight: 800;
            margin: 0 0 .35rem 0;
          }

          .hero-subtitle {
            color: #eafff8 !important;
            font-size: 1rem;
            line-height: 1.5;
            margin: 0;
          }

          .status-grid {
            display: grid;
            grid-template-columns: repeat(3, minmax(0, 1fr));
            gap: .75rem;
            margin: .75rem 0 1rem;
          }

          .status-card {
            background: #ffffff;
            border: 1px solid var(--line);
            border-radius: 8px;
            padding: 1rem;
            box-shadow: 0 2px 8px rgba(15, 23, 42, .06);
          }

          .status-label {
            color: var(--muted) !important;
            font-size: .78rem;
            font-weight: 800;
            text-transform: uppercase;
            letter-spacing: 0;
          }

          .status-value {
            color: var(--ink) !important;
            font-size: 1.35rem;
            font-weight: 800;
            margin-top: .2rem;
            overflow-wrap: anywhere;
          }

          .source-card {
            border: 1px solid var(--line);
            border-radius: 8px;
            padding: 1rem;
            background: #ffffff;
            margin-bottom: .75rem;
            box-shadow: 0 2px 8px rgba(15, 23, 42, .05);
          }

          .source-title {
            color: var(--accent-dark) !important;
            font-weight: 800;
            margin-bottom: .4rem;
            font-size: .95rem;
          }

          .source-text {
            color: #1f2937 !important;
            font-size: .95rem;
            line-height: 1.5;
          }

          [data-testid="stChatMessage"] {
            background: #ffffff !important;
            border: 1px solid var(--line);
            border-radius: 8px;
            padding: .6rem .85rem;
            box-shadow: 0 2px 8px rgba(15, 23, 42, .05);
          }

          [data-testid="stChatMessage"] * {
            color: var(--ink) !important;
          }

          [data-testid="stMarkdownContainer"] p,
          [data-testid="stMarkdownContainer"] li,
          [data-testid="stMarkdownContainer"] span,
          [data-testid="stMarkdownContainer"] div {
            color: var(--ink);
            font-size: 1rem;
            line-height: 1.55;
          }

          .stSelectbox div,
          .stSlider div,
          textarea,
          input {
            color: var(--ink) !important;
          }

          div[data-baseweb="select"] > div,
          textarea,
          input {
            background: #ffffff !important;
            border-color: #94a3b8 !important;
          }

          .stButton button {
            border-radius: 6px;
            border: 1px solid #94a3b8;
            background: #ffffff;
            color: var(--ink) !important;
            font-weight: 700;
            min-height: 2.5rem;
          }

          .stButton button:hover {
            border-color: var(--accent);
            color: var(--accent-dark) !important;
            background: #f0fdfa;
          }

          [data-testid="stExpander"] {
            background: #ffffff !important;
            border: 1px solid var(--line);
            border-radius: 8px;
          }

          .stAlert {
            background: #ffffff !important;
            color: var(--ink) !important;
          }

          code {
            color: var(--ink) !important;
            background: #eef2f7 !important;
          }

          @media (max-width: 800px) {
            .status-grid {
              grid-template-columns: 1fr;
            }
            .hero-title {
              font-size: 1.6rem;
            }
          }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_source_cards(sources: list[dict]) -> None:
    """Render retrieved source metadata and preview text below an answer."""
    if not sources:
        st.info("No retrieved sources to display for this response.")
        return

    for source in sources:
        source_file = escape(str(source.get("source_file") or "Unknown source"))
        department = escape(str(source.get("department") or "Unknown"))
        security_level = escape(str(source.get("security_level") or "Unknown"))
        preview = escape((source.get("content") or "")[:700])
        st.markdown(
            f"""
            <div class="source-card">
              <div class="source-title">{source_file} | chunk {source.get("chunk_id")} | {department} | {security_level}</div>
              <div class="source-text">{preview}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )


st.set_page_config(page_title="Enterprise RAG Assistant", layout="wide")
add_css()

# Streamlit executes the whole script from top to bottom on every interaction.
# Cached functions above keep expensive client creation and index inspection
# from repeating on every rerun.
summary = get_index_summary()

# Sidebar controls define either manual access scope or auto-routing behavior.
# The generated OData filter is displayed so learners can see exactly what is
# sent to Azure AI Search.
with st.sidebar:
    st.subheader("Routing")
    routing_mode = st.radio("Mode", ["Auto", "Manual"], index=0)

    st.subheader("Access Scope")
    department_options = ["HR", "IT", "General", "All"]
    department = st.selectbox("Department", department_options, index=0)
    security_level = st.selectbox("Security level", ["internal"], index=0)
    top = st.slider("Retrieved sources", min_value=1, max_value=5, value=3)
    manual_filter_expression = build_access_filter(department, security_level)

    st.caption("Manual server-side filter")
    st.code(manual_filter_expression, language="text")
    if routing_mode == "Auto":
        st.info("Auto mode routes HR questions to HR docs, IT/helpdesk questions to IT docs, mixed questions to All, and generic questions to direct LLM.")

    st.caption("Live fallback")
    st.code(
        f"enabled={live_fallback_enabled()} | sources={','.join(configured_rag_sources()) or 'none'}",
        language="text",
    )

    st.divider()
    st.subheader("Sample Questions")
    for idx, sample in enumerate(SAMPLE_QUESTIONS):
        if st.button(sample, key=f"sample_{idx}", use_container_width=True):
            st.session_state.pending_prompt = sample

    st.divider()
    st.subheader("Production Notes")
    st.caption("For production, place this app behind Microsoft Entra ID and derive filters from user claims or groups.")

# Header and dashboard cards. These are informational only; the actual health
# check is `summary["schema_ok"]` below.
st.markdown(
    """
    <section class="hero">
      <div class="hero-title">Enterprise RAG Assistant</div>
      <p class="hero-subtitle">
        Ask trusted questions across indexed enterprise documents. Answers are grounded in Azure AI Search sources and include citations.
      </p>
    </section>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    f"""
    <div class="status-grid">
      <div class="status-card"><div class="status-label">Search Index</div><div class="status-value">{summary["index_name"]}</div></div>
      <div class="status-card"><div class="status-label">Indexed Chunks</div><div class="status-value">{summary["document_count"]}</div></div>
      <div class="status-card"><div class="status-label">Source Files</div><div class="status-value">{summary["source_count"]}</div></div>
    </div>
    """,
    unsafe_allow_html=True,
)

if not summary["schema_ok"]:
    st.error(f"Index schema mismatch. Missing fields: {summary['missing_fields']}")

# Show the sampled source files so users can confirm which content is currently
# available in the configured Search index.
with st.expander("Indexed source files", expanded=False):
    for source in summary["sources"]:
        st.write(source)

# Store chat history in Streamlit session state. Without this, every Streamlit
# rerun would forget earlier messages.
if "messages" not in st.session_state:
    st.session_state.messages = [
        {
            "role": "assistant",
            "content": "I am ready. Ask about HR policies, employee records, IT security, support rules, benefits, leave, or payroll.",
            "sources": [],
            "elapsed_ms": None,
            "route_note": None,
        }
    ]

# Replay existing messages on every rerun so the chat transcript remains
# visible after widgets trigger a page refresh.
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.write(message["content"])
        if message.get("route_note"):
            st.caption(message["route_note"])
        if message.get("elapsed_ms") is not None:
            st.caption(f"Completed in {message['elapsed_ms']} ms")
        if message.get("sources"):
            with st.expander("Sources", expanded=True):
                render_source_cards(message["sources"])

# Sidebar sample buttons store their prompt in session state. The chat input
# stores typed text separately. This line gives typed text priority when both
# exist.
prompt = st.session_state.pop("pending_prompt", None) if "pending_prompt" in st.session_state else None
typed_prompt = st.chat_input("Ask a question about your indexed enterprise documents")
prompt = typed_prompt or prompt

if prompt:
    # Add and display the user's message immediately, then compute the assistant
    # response in the same rerun.
    st.session_state.messages.append(
        {"role": "user", "content": prompt, "sources": [], "elapsed_ms": None, "route_note": None}
    )
    with st.chat_message("user"):
        st.write(prompt)

    with st.chat_message("assistant"):
        started = perf_counter()
        if is_greeting(prompt):
            # Greetings do not need retrieval or model calls.
            result = {
                "answer": "Hi! Ask me a business question about the indexed enterprise documents, and I will answer with citations.",
                "sources": [],
            }
            route_note = "Route: greeting | No retrieval"
        else:
            # Auto mode uses transparent keyword routing. Manual mode always
            # forces enterprise document retrieval with the selected scope.
            route = route_query(prompt) if routing_mode == "Auto" else {
                "route": "manual",
                "department": department,
                "retrieval": True,
                "tool": None,
                "reason": f"Manual mode selected. Searching {department} scope.",
            }
            active_filter_expression = build_access_filter(route["department"], security_level)
            if routing_mode == "Manual":
                active_filter_expression = manual_filter_expression

            # This note is shown under the answer so the user can inspect why a
            # particular route/filter was selected.
            route_note = (
                f"Route: {route['route']} | {route['reason']} "
                f"| Filter: {active_filter_expression if route['retrieval'] else 'none'}"
            )

            spinner_text = (
                f"Calling {route['tool']} tool..."
                if route.get("tool")
                else "Answering directly with the LLM..."
                if not route["retrieval"]
                else f"Retrieving {route['department']} sources and generating grounded answer..."
            )
            # Pick the function that implements the selected route.
            with st.spinner(spinner_text):
                try:
                    if route.get("tool") == "weather":
                        result = answer_weather_question(prompt)
                    elif route.get("tool") == "sports":
                        result = answer_sports_question(prompt)
                    elif route.get("tool") == "news":
                        result = answer_news_question(prompt)
                    elif route["retrieval"]:
                        result = answer_question(prompt, active_filter_expression, top)
                    else:
                        result = answer_generic_question(prompt)
                except Exception as exc:
                    # Surface Azure/API/configuration failures directly in the
                    # UI so demo troubleshooting is immediate.
                    st.error(str(exc))
                    st.stop()

        elapsed_ms = round((perf_counter() - started) * 1000)
        # Render answer, route note, timing, and source cards for the current
        # prompt before persisting the assistant message.
        st.write(result["answer"])
        st.caption(route_note)
        st.caption(f"Completed in {elapsed_ms} ms")
        if result["sources"]:
            with st.expander("Sources", expanded=True):
                render_source_cards(result["sources"])

    # Persist the assistant response so it is replayed on the next Streamlit
    # rerun.
    st.session_state.messages.append(
        {
            "role": "assistant",
            "content": result["answer"],
            "sources": result["sources"],
            "elapsed_ms": elapsed_ms,
            "route_note": route_note,
        }
    )
